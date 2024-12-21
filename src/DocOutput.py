from docx import Document
from latex2word import LatexToWordElement

# 这里的各类表达式，我认为应该用一个字典
LLC_EXP = {
    'Gmax': r'G_{max}=n_{p2s}\frac{V_{omax}}{V_{inmin}}',
    'Gmin': r'G_{min}=n_{p2s}\frac{V_{omin}}{V_{inmax}}',
}

# 新建一个word文档
doc = Document()
paragraph = doc.add_paragraph()

LatexToWordElement(LLC_EXP['Gmax']).add_latex_to_paragraph(paragraph)
# paragraph.add_run(f'\n\n')
LatexToWordElement(LLC_EXP['Gmin']).add_latex_to_paragraph(paragraph)

# 输出word文档
doc.save('latex.docx')